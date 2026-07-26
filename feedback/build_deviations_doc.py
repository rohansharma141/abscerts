# -*- coding: utf-8 -*-
"""Generate a Word doc tabulating where the new ABS test site deviates from live abscerts.com."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x08, 0x1F, 0x4D)

doc = Document()

# --- landscape ---
sec = doc.sections[0]
sec.orientation = WD_ORIENT.LANDSCAPE
sec.page_width, sec.page_height = sec.page_height, sec.page_width
sec.left_margin = sec.right_margin = Inches(0.5)
sec.top_margin = sec.bottom_margin = Inches(0.6)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(9)

# --- title ---
t = doc.add_paragraph()
r = t.add_run('ABS Certifications & Advisory — Content Positioning Changes')
r.bold = True
r.font.size = Pt(16)
r.font.color.rgb = NAVY
sub = doc.add_paragraph()
r = sub.add_run('Where the new test site (abscerts.pages.dev) deviates from the live site (abscerts.com), and why')
r.italic = True
r.font.size = Pt(10.5)
r.font.color.rgb = RGBColor(0x55,0x55,0x55)

intro = doc.add_paragraph()
intro.add_run(
    'This table lists only the substantive differences. It compares three sources: (1) the current live '
    'abscerts.com; (2) the independent content review supplied by the client; and (3) what the new test site now '
    'does. The live site is itself internally contradictory — it calls ABS an "independent certification body" that '
    '"provides ISO certification" while elsewhere describing consulting and advisory work; the review resolves this in '
    'favour of the consulting positioning. In the final column, '
).font.size = Pt(9)
b = intro.add_run('(Implemented)'); b.bold=True; b.font.size=Pt(9)
intro.add_run(' = already live on the test site; ').font.size=Pt(9)
b = intro.add_run('(Planned)'); b.bold=True; b.font.size=Pt(9)
intro.add_run(' = agreed direction, pending a decision or client-supplied facts.').font.size=Pt(9)

# --- rows: (section, live, review, test) ---
rows = [
 ("Overall positioning / self-description",
  'Calls ABS an "independent certification body which provides ISO certification to organisations across the globe" (footer + About), while elsewhere also calling itself advisory/consulting.',
  'Reposition as a consulting, implementation, training and certification-readiness firm. Certification audits, decisions and certificate issuance are performed independently by the client’s selected accredited certification body.',
  'ABS described throughout as "a consulting and advisory firm that helps organisations achieve ISO, SOC 2 and CMMi certification." (Implemented)'),

 ("Accreditation claims (IAS / IAF / MLA)",
  'Makes NO accreditation claim anywhere on the live site.',
  'Remove accreditation badges/wording unless they belong to a separately named certification partner whose use is formally authorised.',
  'The earlier test build (ported from the v7 mockup) had ADDED "IAS/IAF accredited" in ~40 places; all removed. Now aligned with the live site’s silence. (Implemented)'),

 ("Who issues the certificate",
  'Implies ABS itself provides / issues ISO certification.',
  'The independent certification body conducts the audit, makes the certification decision and issues the certificate.',
  'States the certificate is issued by an independent certification body; ABS prepares and supports the client. (Implemented)'),

 ("Certification audits (Stage 1 / Stage 2 / surveillance)",
  'Silent — does not describe ABS conducting audits.',
  'Stage 1, Stage 2 and surveillance audits are conducted by the appointed certification body; ABS prepares the organisation but does not make the decision.',
  'Educational explanation retained but reframed — "your certification body carries out the audit; we prepare your team and evidence." (Implemented)'),

 ("Auditors vs consultants",
  'Silent — no "our auditors" / "lead auditors" language.',
  'Replace auditor language with consultants; refer to "the appointed certification body’s auditors" for the formal audit.',
  '"lead auditors" changed to "consultants" throughout (hero, regional cards, why-us, stats, service prose). (Implemented)'),

 ("About page",
  'The most certification-body-heavy content — repeatedly identifies ABS as a certification body and attributes IAS/IAF accreditation directly to ABS.',
  'Replace the accreditation section with an "Our role and independence" section: ABS does not make certification decisions or issue accredited certificates.',
  'About page rewritten; accreditation section replaced with "How we work"; explicitly states "we do not issue certificates ourselves." (Implemented)'),

 ("Process / “How we work” page",
  'Silent — no defined ABS process.',
  'Cited as the strongest page — keep as the model. Four phases: scoping, gap analysis, implementation & training, then an audit by the client’s certification body.',
  'Rewritten to exactly this; opens with a "Who does what" section stating the consultant / certification-body split. (Implemented)'),

 ("Footer disclaimer",
  'No disclaimer; footer calls ABS a certification body.',
  'Add an explicit disclaimer that ABS does not issue accredited certificates or SOC attestation reports; formal decisions are made independently by the authorised third party.',
  'Footer reworded to consulting/advisory; full formal disclaimer paragraph still to be added. (Planned)'),

 ("Statistic — “certificates issued”",
  'General claim that ABS "provides ISO certification globally"; no specific number.',
  '"[Verified number]+ organisations supported." Do not count certificates issued by certification bodies as issued by ABS.',
  'Still displays "1,200+ certificates issued… since 2008." Awaiting a client-verified "organisations supported" figure. (Planned — needs client fact)'),

 ("Statistic — “pass rate”",
  'None.',
  'Replace "98% first-time pass rate" with a verified consulting metric, or remove — a consultant should not imply control over the certification decision.',
  'Still displays "98% first-time pass rate." Awaiting client input. (Planned — needs client fact)'),

 ("Years in business",
  'ABS positioned as established; live site does not state a consistent founding year on the reviewed pages.',
  'Use "18 years" / "since 2004" only where documentary evidence exists.',
  'Test site currently states "Est. 2008." The founding date is inconsistent across sources and needs the client to confirm one figure. (Planned — needs client fact)'),

 ("Navigation & category names",
  'Nav uses "ISO Certifications", "Cyber Security", "Industry & Food", "Personnel Certifications".',
  'Rename to consulting equivalents, e.g. "ISO Management System Consulting", "Cybersecurity & Compliance Advisory"; remove/rename Personnel Certifications.',
  'Nav labels currently unchanged (retains search-friendly terms). Full renames pending a decision on the positioning-vs-SEO trade-off. (Planned — your decision)'),

 ("Homepage hero heading",
  'Certification-led headline.',
  '"Compliance systems that strengthen your business and prepare you for certification."',
  'Currently "Certification that opens doors across the world" with consulting sub-copy. Heading rewrite pending. (Planned — your decision)'),

 ("SOC 2 page",
  'Presents SOC 2 as a certificate ("timeline to certificate", accreditation badge).',
  'SOC 2 is an attestation report issued by a licensed CPA firm — use "readiness period"; badge → "AICPA Trust Services Criteria readiness".',
  'SOC 2 page still uses certificate/"timeline to certificate" wording. Correction queued as a factual fix. (Planned)'),

 ("PCI DSS page",
  'Uses certificate / "timeline to certificate" framing.',
  'PCI DSS validation yields an SAQ, Attestation of Compliance or Report on Compliance via an authorised QSA — not a certificate.',
  'PCI DSS page correction queued as a factual fix. (Planned)'),

 ("HACCP page",
  'Titled "HACCP Certification"; describes ABS certifying the client’s HACCP system.',
  'Retitle "HACCP Implementation & Certification Readiness"; the certificate is provided by the appointed certification organisation.',
  'HACCP page retitle/reframe queued as a factual fix. (Planned)'),

 ("CMMI terminology",
  'Uses "CMMi"; references SCAMPI.',
  'Standardise to "CMMI"; replace SCAMPI with the current CMMI Appraisal Method; formal appraisals by authorised appraisal professionals.',
  'Terminology standardisation queued as a factual fix. (Planned)'),

 ("Client testimonials",
  'Real client testimonials referring to "ABS auditors", "certification provider", "certification services".',
  'Reword to client-approved language separating consulting support from the independent certification audit.',
  'The real testimonials the client supplied are retained verbatim. We will not alter a real client’s quote without their approval. (Planned — needs client approval)'),

 ("Personnel Certifications page",
  'Generic "independent certification body" line; implies ABS certifies individuals.',
  'Should not remain live in current form — remove from nav + de-index, or rename "Professional Competence Development" / "Lead Auditor Registration & Career Guidance".',
  'Page currently live in navigation. Recommended action: pull from nav + noindex now, rebuild once the client confirms whether an approved scheme exists. (Planned — your decision)'),
]

headers = ["S. No.", "Section", "Live site (abscerts.com) — content / intent",
           "Review document — content / intent", "New ABS test site — content / intent"]

table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
table.autofit = False
widths = [Inches(0.45), Inches(1.9), Inches(2.85), Inches(2.85), Inches(2.9)]

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:color'),'auto'); sh.set(qn('w:fill'),hexcolor)
    tcPr.append(sh)

def set_cell(cell, text, width, bold=False, white=False, size=9):
    cell.width = width
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(2)
    run = p.add_run(text)
    run.bold = bold; run.font.size = Pt(size)
    if white: run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

# header row
hdr = table.rows[0].cells
for i,h in enumerate(headers):
    set_cell(hdr[i], h, widths[i], bold=True, white=True, size=9.5)
    shade(hdr[i], '081F4D')

# body
for n,(section,live,review,test) in enumerate(rows, 1):
    cells = table.add_row().cells
    set_cell(cells[0], str(n), widths[0])
    set_cell(cells[1], section, widths[1], bold=True)
    set_cell(cells[2], live, widths[2])
    set_cell(cells[3], review, widths[3])
    set_cell(cells[4], test, widths[4])
    if n % 2 == 0:
        for c in cells: shade(c, 'EEF2F8')

# repeat header row on each page
trPr = table.rows[0]._tr.get_or_add_trPr()
th = OxmlElement('w:tblHeader'); th.set(qn('w:val'),'true'); trPr.append(th)

doc.add_paragraph()
foot = doc.add_paragraph()
fr = foot.add_run('Prepared for internal review. "Live site" reflects abscerts.com as reviewed on 2026-07-20. '
                  'Note: the live site itself still carries the "independent certification body" wording and needs the same correction at source.')
fr.italic = True; fr.font.size = Pt(8); fr.font.color.rgb = RGBColor(0x77,0x77,0x77)

out = 'ABS_deviations_from_live_site.docx'
doc.save(out)
print('SAVED', out)
