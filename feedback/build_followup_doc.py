# -*- coding: utf-8 -*-
"""Generate the ABS follow-up questions & confirmations Word doc (plain-language version)."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x08, 0x1F, 0x4D)
doc = Document()
sec = doc.sections[0]
sec.orientation = WD_ORIENT.LANDSCAPE
sec.page_width, sec.page_height = sec.page_height, sec.page_width
sec.left_margin = sec.right_margin = Inches(0.5)
sec.top_margin = sec.bottom_margin = Inches(0.55)
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(9)

t = doc.add_paragraph(); r = t.add_run('ABS Website — Questions and Confirmations We Need From You')
r.bold = True; r.font.size = Pt(16); r.font.color.rgb = NAVY

intro = doc.add_paragraph()
intro.add_run(
 'What this is: We have already made the main changes from your review. The website now describes ABS as a '
 'consulting and advisory firm (not a certification body), and we have updated the ISO standards to their latest '
 'versions. To finish the rest of the work, we need a few answers, confirmations and materials from you.\n'
).font.size = Pt(9.5)
intro.add_run(
 'How to use this: go through the list below and write your answer in the last column. Each row explains what we '
 'are asking and why. A few rows are marked ').font.size = Pt(9.5)
b = intro.add_run('“needed before we can build this”'); b.bold = True; b.font.size = Pt(9.5)
intro.add_run(' — we cannot finish those parts until you reply.').font.size = Pt(9.5)

SECTIONS = [
 ("A.  Please confirm we got these right", [
  ("CMMI — who runs the appraisal",
   "Please confirm: ABS helps clients get ready for a CMMI appraisal, but does not carry out the official appraisal itself. Is that correct? If ABS does have its own qualified CMMI appraiser and can run official appraisals, please tell us and we will word it that way instead.",
   "The old site said ABS had run CMMI appraisals “since 1991” with its own lead appraisers. Your review said to describe ABS as helping clients prepare, with the official appraisal done by an authorised appraiser. We made that change and want to be sure it is accurate."),
  ("Are the numbers on the site true?",
   "Please confirm each of these figures is true and you have records to back it up: “1,200+ engagements supported”, the first-time success rate, “40+ countries”, and “since 2008”. If any figure is wrong or cannot be backed up, tell us the correct number or ask us to remove it.",
   "The site now shows these figures with a note, “Figures updated July 2026”. Adding a date makes them look officially checked, so we want to be sure they are correct before the site goes public."),
  ("Which start year should we show?",
   "Please tell us the single, correct history to show. For example: “Training since 2004; ABS Certifications & Advisory since 2008.”",
   "Different pages currently suggest different start years — 2008 (the company), 2004 (the training academy) and 1991 (old CMMI text). We want one consistent story across the whole site."),
  ("Wording of the “1,200+” figure",
   "We changed “1,200+ certificates issued” to “1,200+ certification and compliance engagements supported”. Please confirm this wording is fine.",
   "The old wording made it sound like ABS issues the certificates itself. The new wording says ABS supported the projects, which fits the consulting positioning."),
 ]),
 ("B.  Questions about specific standards and claims", [
  ("ISO 27017 — which year to show  (needed before we can build this)",
   "Has ISO officially published the 2026 version of ISO/IEC 27017 yet? If yes, we will update the page to “2026”. If it is not published yet, we will keep it as “2015”.",
   "As of April 2026 the 2026 version was still a final draft, not officially published. We updated every other standard to its new version but held this one until we know it is official."),
  ("Does ABS certify individual people?  (needed before we can build this)",
   "Does ABS run an official, approved scheme that certifies individual people (for example, certifying someone as a qualified lead auditor)? Please answer Yes or No. If Yes, please name the scheme. If No, we will change these pages to say ABS provides training and guidance, and that a training certificate is not the same as official personnel certification.",
   "The “Personnel Certifications” pages currently say ABS certifies a person’s competence. Your review said to only claim this where there is an approved scheme. We have not changed these pages yet because it depends on your answer."),
  ("Inspection services — accreditation",
   "For third-party inspection work, is ABS itself accredited to do it, or does ABS use accredited partner inspectors? We will add a short line to match.",
   "Your review asked us to be clear about accreditation status for inspections, and to mention partners where they are used."),
  ("PCI DSS — ABS’s role",
   "Please confirm: ABS helps clients get ready for PCI DSS, and the official assessment (where one is needed) is carried out by an authorised assessor (a “QSA”), not by ABS.",
   "This matches your review; we have written it this way and want to confirm."),
  ("HIPAA and GDPR wording",
   "Please confirm this is right: there is no official HIPAA “certificate” and no single GDPR “certificate” — ABS provides assessments and readiness support instead.",
   "This matches your review and is already on the pages."),
 ]),
 ("C.  Information and materials we need from you", [
  ("Company legal details  (needed before we can build this)",
   "Please send: the full registered company name, registered office address, company registration number, GST number, an official email address, a contact for privacy questions, and the country whose laws apply.",
   "We need these for the website footer, the Privacy Policy and the Terms page. The Privacy Policy currently still has placeholder text — and even shows the wrong website name, “hmlcerts.com”, which must be replaced."),
  ("Which online booking tool?",
   "For the “Book a call” button, which tool would you like — Calendly, Zoho Bookings, or Microsoft Bookings? Please share the account or booking link.",
   "Right now the button just opens the contact page. Your review asked for a real booking calendar so visitors can pick a time slot themselves."),
  ("Team profiles for the About page",
   "For each senior consultant or leader you want featured, please send: name, role, years of experience, qualifications, any lead-auditor credentials, and a LinkedIn link.",
   "Your review asked for a “Leadership & Experts” section on the About page. Real, named experts build trust with buyers and help the site rank on Google."),
  ("Client logos",
   "Which clients have given written permission to show their logo on the site? Please send those logo files.",
   "We only show client logos where there is written permission. Until then, the logo strip stays hidden."),
  ("Case studies",
   "Please provide 2–3 real, approved client stories. For each: the client, the challenge, what was in scope, what ABS did, what was delivered, how long it took, the outcome, and a short client quote.",
   "The case-studies section is built but empty. Real stories are strong proof for prospects."),
  ("Names on testimonials",
   "For the client testimonials, may we show the person’s name, job title, company, country and service (with their written permission)? Where a client must stay anonymous, is a description such as “CISO, UK-based SaaS company” acceptable?",
   "We have three real testimonials but with very little detail. Adding names (with permission) makes them far more believable."),
  ("Blog authors and review dates",
   "Please give a real author name and job title (and a reviewer name) to show on articles. Also, please tell us which articles have actually been checked or updated recently, so we only add an “updated” note where it is true.",
   "Articles currently show “ABS Certifications” as the author. Real, named authors improve credibility and Google trust. We do not want to claim an article was updated if it was not."),
  ("Third phone number",
   "Should we add the third phone number shown on your current website: +91 99116 02258?",
   "The contact page currently lists two numbers."),
  ("Social media links",
   "Please send the official links for LinkedIn, X (Twitter) and YouTube.",
   "The footer social icons are placeholders at the moment and do not link anywhere."),
  ("Information security (ISMS) policy",
   "Do you have an approved Information Security policy statement we can publish? If not, we will remove the “ISMS policy” link for now.",
   "The footer shows an “ISMS policy” link, but there is no page behind it yet."),
 ]),
 ("D.  Decisions about the site’s structure", [
  ("Top menu layout",
   "Would you like the top menu changed to: Services · Industries · Standards · Resources · About · Contact (with a drop-down for the standards)? Your suggested list did not include “Process” or “Blog” — since you praised the Process page, where would you like Process and Blog to sit?",
   "The current menu is: Home · About · Services · Process · Blog · Contact."),
  ("Grouping HR, Data and Agile",
   "Should we group the HR, Data Analytics and Agile services under a separate “Business Advisory” heading, as your review suggested?",
   "Right now all ten service areas sit at the same level in the menu."),
  ("Service page titles",
   "Would you like service pages renamed to the consulting style — for example “ISO 9001 Certification Consulting”, “ISO 27001 Implementation & Certification Support”? Please confirm you want this pattern.",
   "This is a rename across roughly 70 service pages, so we want to confirm before doing it."),
  ("Web addresses of pages",
   "Your review suggested short web addresses like /iso-27001-consulting. Ours are currently /services/iso-27001. Because the new site is not live yet, we can set the final addresses now. Which do you prefer? (We will set up redirects either way.)",
   "Changing web addresses after launch needs redirects and can affect Google ranking, so it is best decided before the site goes live."),
  ("Country pages",
   "Which countries should get their own page (for example UAE, USA, UK)? To do these well we will need genuinely local details from you for each country — thin or copy-pasted country pages can actually hurt Google ranking.",
   "There are no country-specific pages yet."),
  ("Industry pages",
   "We have 4 industry pages (finance, healthcare, manufacturing, and software/SaaS). Would you like more — for example construction, logistics, education, IT services? If so, please share the key points for each new sector.",
   "Your review suggested covering more industries."),
  ("Contact form",
   "Would you like a shorter, two-step enquiry form, and the option for visitors to upload a file (such as an existing certificate or a tender document)? Any limits on file type or size we should set?",
   "The current form is a single page with no file upload."),
 ]),
 ("E.  Legal review", [
  ("Legal sign-off on policy pages  (needed before we can publish)",
   "Will your legal team write or approve the final Privacy Policy, Cookie Policy and Terms? Also, is a data-retention period of “up to 3 years” acceptable to you?",
   "We will draft these pages in a clear structure based on your notes, but the exact legal wording should be checked by a lawyer before it goes public."),
 ]),
]

headers = ["No.", "Topic", "What we are asking", "Why we are asking / where things stand", "Your answer"]
widths = [Inches(0.4), Inches(2.0), Inches(3.5), Inches(2.9), Inches(2.2)]
table = doc.add_table(rows=1, cols=5); table.style = 'Table Grid'; table.autofit = False

def shade(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr(); sh = OxmlElement('w:shd')
    sh.set(qn('w:val'),'clear'); sh.set(qn('w:color'),'auto'); sh.set(qn('w:fill'),hexc); tcPr.append(sh)
def setc(cell, text, w, bold=False, white=False, size=9):
    cell.width = w; p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(2)
    run = p.add_run(text); run.bold = bold; run.font.size = Pt(size)
    if white: run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

hdr = table.rows[0].cells
for i,h in enumerate(headers):
    setc(hdr[i], h, widths[i], bold=True, white=True, size=9.5); shade(hdr[i], '081F4D')
trPr = table.rows[0]._tr.get_or_add_trPr(); th=OxmlElement('w:tblHeader'); th.set(qn('w:val'),'true'); trPr.append(th)

n = 0
for section, items in SECTIONS:
    row = table.add_row().cells
    merged = row[0].merge(row[1]).merge(row[2]).merge(row[3]).merge(row[4])
    setc(merged, section, widths[0], bold=True, size=11)
    shade(merged, 'D8E0EE')
    for topic, need, ctx in items:
        n += 1
        c = table.add_row().cells
        setc(c[0], str(n), widths[0]); setc(c[1], topic, widths[1], bold=True)
        setc(c[2], need, widths[2]); setc(c[3], ctx, widths[3]); setc(c[4], "", widths[4])

doc.add_paragraph()
f = doc.add_paragraph(); fr = f.add_run(
 'Thank you. Once we have your answers to sections A–C we can finish the website text. Section D decides how much '
 'work the next stage involves, and Section E is about the legal pages.')
fr.italic = True; fr.font.size = Pt(8.5); fr.font.color.rgb = RGBColor(0x66,0x66,0x66)

doc.save('ABS_followup_questions.docx'); print('SAVED ABS_followup_questions.docx  | items:', n)
