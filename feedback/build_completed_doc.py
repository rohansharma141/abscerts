# -*- coding: utf-8 -*-
"""Generate the 'what we have already changed' Word doc (plain language, confirmation column)."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x08, 0x1F, 0x4D)
doc = Document()
sec = doc.sections[0]
sec.orientation = WD_ORIENT.LANDSCAPE
sec.page_width, sec.page_height = sec.page_height, sec.page_width
sec.left_margin = sec.right_margin = Inches(0.5)
sec.top_margin = sec.bottom_margin = Inches(0.55)
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(9)

t = doc.add_paragraph(); r = t.add_run('ABS Website — What We Have Already Changed')
r.bold = True; r.font.size = Pt(16); r.font.color.rgb = NAVY

intro = doc.add_paragraph()
intro.add_run(
 'This is a summary of the changes we have already made to the website, based on your review. The main goal was to '
 'describe ABS correctly across the whole site — as a consulting and advisory firm that helps clients get ready for '
 'certification, with the official certificate, report or decision made by an independent third party (a certification '
 'body, a licensed CPA firm, a QSA, or an authorised appraiser).\n'
).font.size = Pt(9.5)
intro.add_run(
 'Everything below is already live on the test site. Please read through and, in the last column, confirm each change '
 'looks right — or add a short note where you would like something different. Anything we still need from you to finish '
 'the remaining work is listed in the separate questions document.').font.size = Pt(9.5)

SECTIONS = [
 ("A.  The main change — how ABS is described", [
  ("Whole-site positioning",
   "The entire website now describes ABS as a consulting and advisory firm — not a certification body. It explains that ABS helps clients implement and get ready, and that the official certificate, report or decision is made by an independent third party.",
   "Your review said the old site mixed the two roles and sometimes made ABS sound like it issues certificates itself."),
  ("Removed “IAS / IAF accredited” claims",
   "Removed all “IAS / IAF accredited” wording. It had appeared in about 40 places (badges, homepage, service pages, footer).",
   "These claims were not on your live website — they came from the original design template. Your review asked to remove accreditation claims unless they belong to a named, authorised partner."),
  ("Who issues the certificate",
   "Every relevant page now says clearly that the certificate is issued by an independent certification body, not by ABS.",
   "To make the split between ABS’s role and the certifier’s role obvious."),
  ("“Auditors” changed to “consultants”",
   "Wording that called ABS’s people “auditors” or “lead auditors” has been changed to “consultants”. The official audit is described as being done by the certification body.",
   "So the site does not suggest ABS carries out the official certification audit."),
 ]),
 ("B.  Key pages rewritten", [
  ("About page",
   "Rewrote the About page: ABS is described as an independent consulting firm; added an “Our role and independence” note explaining the certification decision is made by the independent body; clarified that work is delivered by internal consultants and approved associates.",
   "The old About page most strongly implied ABS was a certification body."),
  ("“How we work” (Process) page",
   "Rewrote the process. It now clearly shows ABS preparing the client, and the certification body carrying out Stage 1 and Stage 2 and issuing the certificate. We also softened the wording so it does not promise the client will pass.",
   "Your review pointed to this page as the model to follow, and asked us not to imply a guaranteed outcome."),
  ("Homepage",
   "New main heading: “Global ISO, SOC 2 and compliance advisory services.” The main button now says “Request a fixed-price proposal.” The statistics were reworded (for example “1,200+ certification and compliance engagements supported”), and we added a small note, “Figures updated July 2026.”",
   "To lead with the consulting message and stop the numbers implying ABS issues certificates. (We have asked you to confirm the figures are accurate.)"),
  ("Contact page",
   "Changed the response-time promise to “within one working day, Monday–Friday, excluding public holidays.” Reworded the consent checkbox and added a separate optional box for marketing emails.",
   "A clearer, more realistic promise, and cleaner consent wording."),
  ("Footer",
   "Added a standard disclaimer: ABS does not make certification or accreditation decisions; those are made by the independent third party; and engaging ABS does not guarantee a certificate. The copyright year now updates automatically (currently “© 2008–2026”).",
   "To set clear expectations and keep the year current."),
 ]),
 ("C.  Specific services and standards", [
  ("CMMI",
   "Removed the old “since 1991 / our lead appraisers run appraisals” claims. CMMI is now described as ABS helping clients get ready, with the official appraisal done by an authorised appraiser. We also corrected the spelling “CMMi” to “CMMI” and replaced the outdated “SCAMPI” term with the current “CMMI Appraisal Method.”",
   "Your review repositioned CMMI as readiness support. (Please confirm ABS does not run official appraisals itself — this is in the questions document.)"),
  ("SOC 2",
   "Clarified that SOC 2 results in an attestation report from an independent licensed CPA firm — not a certificate from ABS.",
   "SOC 2 is often wrongly called a “certificate”; your review asked to correct this."),
  ("PCI DSS",
   "Clarified that the outcome is an SAQ, Attestation of Compliance or Report on Compliance through an authorised QSA — not a certificate — and that ABS provides readiness support.",
   "To describe PCI DSS accurately."),
  ("HACCP",
   "Renamed the page to “HACCP Implementation & Certification Readiness” and made clear the certificate is provided by the certification body.",
   "The old page implied ABS certifies the HACCP system itself."),
  ("ISO 27017 / 27018 / 27701",
   "Reworded these so they are described as guidance used alongside ISO 27001 — or, for the new ISO 27701:2025, as a standard that can stand on its own — rather than separate certificates from ABS.",
   "Your review said these are not standalone certifications."),
  ("HIPAA / GDPR",
   "Made clear there is no official HIPAA certificate and no single GDPR certificate; ABS provides assessments and readiness support.",
   "To avoid implying a certificate exists where it does not."),
 ]),
 ("D.  Standard versions brought up to date", [
  ("Latest editions applied",
   "Updated the standard versions across the site: ISO/IEC 27701 → 2025 (now standalone); ISO/IEC 27018 → 2025; ISO 55001 → 2024; ISO 37001 → 2025 (with transition information; deadline Feb 2027); ISO 21001 → 2025; ISO 14001 → added the 2026 transition information (published Apr 2026; deadline Apr 2029).",
   "Your review listed the current editions; we checked each one against ISO’s official catalogue and they were correct."),
  ("Withdrawn standards handled",
   "ISO 19600 (withdrawn) — created a new page for its replacement, ISO 37301:2021, and set the old web address to redirect there. ISO 29990 (withdrawn) — removed it and set its old address to redirect to ISO 21001.",
   "These standards no longer exist, so keeping them as-is would be inaccurate."),
  ("ISO 27017 held for now",
   "We kept ISO 27017 at the 2015 version for the moment.",
   "The 2026 version appeared to still be a final draft, not officially published. We have asked you to confirm before we change it."),
 ]),
 ("E.  FAQ and fees", [
  ("New FAQ answers",
   "Added clear answers: “Does ABS issue ISO certificates?” (no — the certification body does); how IAF accreditation actually works and how to check a certification body; “are all ISO 27001 Annex A controls required?” (no); and “are ISO 27017/27018 standalone certifications?” (no).",
   "Your review asked for these clarifications."),
  ("Fees explained clearly",
   "Made clear that ABS’s consulting fees are separate from the certification body’s own audit and surveillance fees. We also removed an example price that mixed the two together.",
   "So clients understand what they pay ABS versus what they pay the certifier."),
 ]),
]

headers = ["No.", "Area", "What we changed", "Why we changed it", "Looks right? / your note"]
widths = [Inches(0.4), Inches(1.9), Inches(3.6), Inches(2.9), Inches(2.2)]
table = doc.add_table(rows=1, cols=5); table.style = 'Table Grid'; table.autofit = False

def shade(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr(); sh = OxmlElement('w:shd')
    sh.set(qn('w:val'),'clear'); sh.set(qn('w:color'),'auto'); sh.set(qn('w:fill'),hexc); tcPr.append(sh)
def setc(cell, text, w, bold=False, white=False, size=9):
    cell.width = w; p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(2)
    run = p.add_run(text); run.bold = bold; run.font.size = Pt(size)
    if white: run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

hdr = table.rows[0].cells
for i,h in enumerate(headers):
    setc(hdr[i], h, widths[i], bold=True, white=True, size=9.5); shade(hdr[i], '081F4D')
trPr = table.rows[0]._tr.get_or_add_trPr(); th=OxmlElement('w:tblHeader'); th.set(qn('w:val'),'true'); trPr.append(th)

n = 0
for section, items in SECTIONS:
    row = table.add_row().cells
    merged = row[0].merge(row[1]).merge(row[2]).merge(row[3]).merge(row[4])
    setc(merged, section, widths[0], bold=True, size=11); shade(merged, 'D8E0EE')
    for area, what, why in items:
        n += 1
        c = table.add_row().cells
        setc(c[0], str(n), widths[0]); setc(c[1], area, widths[1], bold=True)
        setc(c[2], what, widths[2]); setc(c[3], why, widths[3]); setc(c[4], "", widths[4])

doc.add_paragraph()
f = doc.add_paragraph(); fr = f.add_run(
 'Everything above is already live on the test site. A simple “looks right” against each row is all we need; please note '
 'anything you would like changed. What we still need from you to finish the rest is in the separate questions document.')
fr.italic = True; fr.font.size = Pt(8.5); fr.font.color.rgb = RGBColor(0x66,0x66,0x66)

doc.save('ABS_completed_changes.docx'); print('SAVED ABS_completed_changes.docx  | items:', n)
